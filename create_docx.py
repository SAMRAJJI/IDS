from docx import Document
from docx.shared import Inches
import os

# Create a new Document
doc = Document()

# Title
doc.add_heading('Intrusion Detection System with Deep Learning - Project Documentation', 0)

# Introduction
doc.add_heading('1. Introduction', level=1)
doc.add_paragraph(
    "This project implements an Intrusion Detection System (IDS) using various Deep Learning techniques. "
    "It covers multiple datasets and model architectures to detect network intrusions effectively. "
    "The project includes web interfaces for real-time prediction and comprehensive training pipelines."
)

# Project Structure
doc.add_heading('2. Project Structure', level=1)
doc.add_paragraph(
    "The workspace is organized as follows:\n"
    "- ids/: Implementation using NSL-KDD dataset with simple DNN\n"
    "- IDS with DL - CIC2017/: Implementation using CIC-IDS2017 dataset with AE-DNN model\n"
    "- unified_ids/: Unified Flask application supporting both datasets\n"
    "- SNN.ipynb: Jupyter notebook exploring Spiking Neural Networks for IDS\n"
    "- cic_sample_inspection.ipynb & cic_sample_inspection2.ipynb: Data exploration notebooks\n"
    "- env/: Python virtual environment"
)

# Technologies Used
doc.add_heading('3. Technologies Used', level=1)

doc.add_heading('Programming Language:', level=2)
doc.add_paragraph("Python 3.11.3")

doc.add_heading('Frameworks and Libraries:', level=2)
libraries = [
    "TensorFlow 2.12+ (Deep Learning framework)",
    "Flask (Web framework for APIs)",
    "Pandas (Data manipulation)",
    "NumPy (Numerical computing)",
    "Scikit-learn (Machine learning utilities)",
    "Joblib (Model serialization)",
    "Matplotlib (Plotting)",
    "Seaborn (Statistical visualization)",
    "python-docx (Document generation)",
    "OpenCV (Computer vision for face recognition)",
    "face_recognition (Face detection library)",
    "bcrypt (Password hashing)"
]
for lib in libraries:
    doc.add_paragraph(lib, style='List Bullet')

# Datasets Used
doc.add_heading('4. Datasets Used', level=1)

doc.add_heading('NSL-KDD Dataset:', level=2)
doc.add_paragraph(
    "The NSL-KDD dataset is an improved version of the KDD Cup 1999 dataset. "
    "It contains 41 features extracted from network traffic, including basic features, "
    "content features, and traffic features. The dataset includes various attack types "
    "such as DoS, Probe, R2L, and U2R attacks."
)

doc.add_heading('CIC-IDS2017 Dataset:', level=2)
doc.add_paragraph(
    "The CIC-IDS2017 dataset was created by the Canadian Institute for Cybersecurity. "
    "It contains modern attack scenarios including DoS, DDoS, Brute Force, XSS, SQL Injection, "
    "and infiltration attacks. The dataset has 78 features and covers traffic from Monday to Friday."
)

# Model Architectures
doc.add_heading('5. Model Architectures', level=1)

doc.add_heading('AE-DNN (Attention-Enhanced Deep Neural Network):', level=2)
doc.add_paragraph(
    "The AE-DNN model incorporates attention mechanisms to focus on important features. "
    "Architecture includes:\n"
    "- Dense layers with batch normalization and dropout\n"
    "- Multi-head attention layers\n"
    "- Residual connections\n"
    "- Output layer with sigmoid activation for binary classification"
)

doc.add_heading('Simple DNN:', level=2)
doc.add_paragraph(
    "A simpler deep neural network architecture used for baseline comparison. "
    "Consists of multiple dense layers with ReLU activation and dropout regularization."
)

# Training Results
doc.add_heading('6. Training Results', level=1)

doc.add_heading('NSL-KDD Model Performance:', level=2)
doc.add_paragraph(
    "After 33 epochs of training:\n"
    "- Final Training Accuracy: 99.08%\n"
    "- Final Validation Accuracy: 99.05%\n"
    "- Final Training Loss: 0.022\n"
    "- Final Validation Loss: 0.028\n"
    "- AUC Score: 0.9994"
)

doc.add_heading('CIC-IDS2017 Model Performance:', level=2)
doc.add_paragraph(
    "After 20 epochs of training:\n"
    "- Final Training Accuracy: 96.80%\n"
    "- Final Validation Accuracy: 96.97%\n"
    "- Final Training Loss: 0.058\n"
    "- Final Validation Loss: 0.055\n"
    "- AUC Score: 0.9975"
)

# Web Applications
doc.add_heading('7. Web Applications', level=1)

doc.add_heading('Individual Dataset Applications:', level=2)
doc.add_paragraph(
    "Separate Flask applications for each dataset:\n"
    "- ids/app.py: NSL-KDD prediction interface\n"
    "- IDS with DL - CIC2017/app.py: CIC-IDS2017 prediction interface\n"
    "Both applications provide web forms for inputting network features and real-time prediction."
)

doc.add_heading('Unified Application:', level=2)
doc.add_paragraph(
    "The unified_ids/ application combines both models in a single interface. "
    "Features include:\n"
    "- Support for both NSL-KDD and CIC-IDS2017 datasets\n"
    "- Face recognition authentication for admin access\n"
    "- Password-based authentication\n"
    "- User management system\n"
    "- Real-time prediction capabilities"
)

# Authentication System
doc.add_heading('8. Authentication System', level=1)
doc.add_paragraph(
    "The unified application includes a multi-modal authentication system:\n"
    "- Face recognition using OpenCV and face_recognition library\n"
    "- Password authentication with bcrypt hashing\n"
    "- Admin setup script for initial face registration\n"
    "- User management with JSON-based storage"
)

# Data Preprocessing
doc.add_heading('9. Data Preprocessing', level=1)
doc.add_paragraph(
    "Comprehensive preprocessing pipelines for both datasets:\n"
    "- Categorical feature encoding using LabelEncoder\n"
    "- Numerical feature scaling using StandardScaler\n"
    "- Handling of missing values and outliers\n"
    "- Feature selection and engineering\n"
    "- Train/validation/test split (70/15/15)"
)

# UI Presets
doc.add_heading('10. UI Presets', level=1)
doc.add_paragraph(
    "Pre-configured input presets for common attack scenarios:\n"
    "- Normal traffic patterns\n"
    "- DoS (Denial of Service) attacks\n"
    "- DDoS (Distributed DoS) attacks\n"
    "- Port scanning attacks\n"
    "- Brute force attacks\n"
    "Presets stored in cic_cic_js_presets.txt and cic_cic_presets.json files"
)

# Future Work
doc.add_heading('11. Future Work', level=1)
doc.add_paragraph(
    "Potential enhancements and extensions:\n"
    "- Implementation of Spiking Neural Networks (SNN) for IDS\n"
    "- Integration of additional datasets\n"
    "- Real-time network traffic monitoring\n"
    "- Ensemble model approaches\n"
    "- Explainable AI techniques for intrusion detection\n"
    "- Deployment on cloud platforms"
)

# How to Run
doc.add_heading('12. How to Run the Project', level=1)

doc.add_heading('Training Models:', level=2)
doc.add_paragraph(
    "For NSL-KDD:\n"
    "cd ids\n"
    "python main.py\n"
    "\n"
    "For CIC-IDS2017:\n"
    "cd IDS with DL - CIC2017\n"
    "python main.py"
)

doc.add_heading('Running Web Applications:', level=2)
doc.add_paragraph(
    "Individual apps:\n"
    "python app.py (from respective directories)\n"
    "\n"
    "Unified app:\n"
    "cd unified_ids\n"
    "python setup_admin.py  # First time only\n"
    "python app.py"
)

# Conclusion
doc.add_heading('13. Conclusion', level=1)
doc.add_paragraph(
    "This project demonstrates a comprehensive approach to intrusion detection using deep learning. "
    "It successfully implements multiple model architectures on different datasets, achieving high accuracy "
    "scores. The web interfaces provide user-friendly access to the prediction capabilities, and the "
    "authentication system ensures secure access. The modular design allows for easy extension and "
    "integration of new datasets and models."
)

# Save the document
doc.save('IDS_Project_Documentation.docx')
print("Documentation saved as 'IDS_Project_Documentation.docx'")

doc.add_heading('Tools:', level=2)
tools = [
    "Visual Studio Code (IDE)",
    "Jupyter Notebook (Interactive development)",
    "Git (Version control)",
    "Kaggle API (Dataset downloading)"
]
for tool in tools:
    doc.add_paragraph(tool, style='List Bullet')

# Datasets
doc.add_heading('4. Datasets', level=1)

doc.add_heading('NSL-KDD:', level=2)
doc.add_paragraph(
    "Classic intrusion detection dataset with 41 features. "
    "Contains network traffic data labeled as normal or various attack types. "
    "Used in the 'ids/' implementation."
)

doc.add_heading('CIC-IDS2017:', level=2)
doc.add_paragraph(
    "Modern dataset from Canadian Institute for Cybersecurity. "
    "Contains 78 features from real network traffic. "
    "Includes various attack scenarios like DDoS, brute force, etc. "
    "Used in the 'IDS with DL - CIC2017/' implementation."
)

# Models and Architectures
doc.add_heading('5. Models and Architectures', level=1)

doc.add_heading('Simple DNN (NSL-KDD):', level=2)
doc.add_paragraph(
    "A feedforward neural network with the following architecture:\n"
    "- Input layer (41 features)\n"
    "- Dense(128) + BatchNorm + Dropout(0.3)\n"
    "- Dense(64) + BatchNorm + Dropout(0.3)\n"
    "- Dense(32) + Dropout(0.2)\n"
    "- Dense(16)\n"
    "- Output: Sigmoid for binary classification\n"
    "Compiled with Adam optimizer and binary crossentropy loss."
)

doc.add_heading('AE-DNN (CIC-IDS2017):', level=2)
doc.add_paragraph(
    "Attention-Enhanced Deep Neural Network:\n"
    "- Input layer (78 features)\n"
    "- Dense blocks with residual connections\n"
    "- Multi-head attention mechanism\n"
    "- Batch normalization and dropout throughout\n"
    "- Output: Sigmoid for binary classification\n"
    "Includes class weighting to handle imbalanced data."
)

doc.add_heading('Spiking Neural Network (SNN):', level=2)
doc.add_paragraph(
    "Exploratory implementation using neuromorphic computing principles. "
    "Implemented in Jupyter notebook for research purposes."
)

# Implementation Steps
doc.add_heading('6. Implementation Steps', level=1)

steps = [
    "1. Data Collection: Download and prepare NSL-KDD and CIC-IDS2017 datasets",
    "2. Data Preprocessing: Clean data, encode categorical features, scale numerical features",
    "3. Feature Engineering: Select relevant features, handle missing values",
    "4. Model Development: Design and implement neural network architectures",
    "5. Training: Train models with appropriate callbacks (early stopping, checkpoints)",
    "6. Evaluation: Assess model performance using accuracy, precision, recall, AUC",
    "7. Web Interface: Develop Flask applications for real-time prediction",
    "8. Deployment: Create unified interface supporting multiple models"
]

for step in steps:
    doc.add_paragraph(step, style='List Number')

# Training Process
doc.add_heading('7. Training Process', level=1)
doc.add_paragraph(
    "Models are trained using the following methodology:\n"
    "- Data splitting: Train/Validation/Test sets\n"
    "- Class weighting for imbalanced datasets\n"
    "- Early stopping based on validation AUC\n"
    "- Model checkpointing to save best weights\n"
    "- Learning rate scheduling\n"
    "- Batch training with appropriate batch sizes"
)

# Web Interfaces
doc.add_heading('8. Web Interfaces', level=1)

doc.add_heading('Individual Apps:', level=2)
doc.add_paragraph(
    "Separate Flask applications for NSL-KDD and CIC-IDS2017 models, "
    "each providing prediction endpoints and user interfaces."
)

doc.add_heading('Unified App:', level=2)
doc.add_paragraph(
    "A single application that supports both datasets and models, "
    "allowing users to choose the appropriate IDS for their needs."
)

# Evaluation Metrics
doc.add_heading('9. Evaluation Metrics', level=1)
metrics = [
    "Accuracy: Overall correct predictions",
    "Precision: True positives / (True positives + False positives)",
    "Recall: True positives / (True positives + False negatives)",
    "AUC-ROC: Area under the receiver operating characteristic curve"
]
for metric in metrics:
    doc.add_paragraph(metric, style='List Bullet')

# Challenges and Solutions
doc.add_heading('10. Challenges and Solutions', level=1)

challenges = [
    "Imbalanced datasets: Solved using class weighting",
    "High-dimensional data: Handled with appropriate network architectures",
    "Categorical features: Encoded using LabelEncoder",
    "Model serialization: Used TensorFlow/Keras and joblib",
    "Web deployment: Implemented with Flask framework"
]

for challenge in challenges:
    doc.add_paragraph(challenge, style='List Bullet')

# Future Improvements
doc.add_heading('11. Future Improvements', level=1)
improvements = [
    "Implement ensemble methods combining multiple models",
    "Add real-time streaming data processing",
    "Integrate with network monitoring tools",
    "Explore advanced architectures like transformers",
    "Add explainability features (SHAP, LIME)",
    "Deploy on cloud platforms for scalability"
]

for imp in improvements:
    doc.add_paragraph(imp, style='List Bullet')

# Conclusion
doc.add_heading('12. Conclusion', level=1)
doc.add_paragraph(
    "This project demonstrates a comprehensive approach to building IDS using deep learning. "
    "It covers the entire pipeline from data preprocessing to web deployment, "
    "providing both research implementations and practical applications. "
    "The modular design allows for easy extension and integration of new models and datasets."
)

# Save the document
doc.save('IDS_Project_Documentation.docx')
print("Documentation created: IDS_Project_Documentation.docx")